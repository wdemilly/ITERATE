import streamlit as st
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import math
import json
import os
from datetime import datetime

st.set_page_config(page_title="Fiction Writer + Detection Scorer", layout="wide")
st.title("Fiction Chapter Writer")

# ──────────────────────────────────────────────
# SOURCE RHYTHM ANALYSIS
# ──────────────────────────────────────────────

def analyze_source_rhythm(text):
    """
    Analyzes source texts (Dare, Quinn, Faulks, etc.) to extract a
    sentence rhythm fingerprint. This becomes the target for revision,
    not an abstract threshold.
    """
    if not text or not text.strip():
        return None
    
    sentence_endings = re.split(r'(?<=[.!?])\s+(?=[A-Z"\u201C])', text)
    sentences = [s.strip() for s in sentence_endings if s.strip()]
    lengths = [len(s.split()) for s in sentences]
    
    if len(lengths) < 10:
        return None
    
    total = len(lengths)
    mean_len = sum(lengths) / total
    variance = sum((x - mean_len) ** 2 for x in lengths) / total
    std_dev = math.sqrt(variance)
    cv = std_dev / mean_len if mean_len > 0 else 0
    
    ultra_short = sum(1 for l in lengths if l <= 5)
    short = sum(1 for l in lengths if 6 <= l <= 12)
    medium = sum(1 for l in lengths if 13 <= l <= 25)
    long = sum(1 for l in lengths if 26 <= l <= 40)
    very_long = sum(1 for l in lengths if l > 40)
    
    short_after_long = 0
    long_after_short = 0
    same_band = 0
    
    for i in range(1, len(lengths)):
        prev, curr = lengths[i-1], lengths[i]
        if prev >= 25 and curr <= 5:
            short_after_long += 1
        if prev <= 5 and curr >= 25:
            long_after_short += 1
        prev_band = 0 if prev <= 5 else (1 if prev <= 12 else (2 if prev <= 25 else 3))
        curr_band = 0 if curr <= 5 else (1 if curr <= 12 else (2 if curr <= 25 else 3))
        if prev_band == curr_band:
            same_band += 1
    
    transitions = total - 1 if total > 1 else 1
    paratactic = sum(1 for s in sentences if s.count(' and ') >= 3)
    max_len = max(lengths)
    min_len = min(lengths)
    range_ratio = max_len / min_len if min_len > 0 else max_len
    
    return {
        "total_sentences": total,
        "word_count": sum(lengths),
        "mean_length": round(mean_len, 1),
        "cv": round(cv, 3),
        "std_dev": round(std_dev, 1),
        "min_length": min_len,
        "max_length": max_len,
        "range_ratio": round(range_ratio, 1),
        "distribution": {
            "ultra_short_pct": round(100 * ultra_short / total, 1),
            "short_pct": round(100 * short / total, 1),
            "medium_pct": round(100 * medium / total, 1),
            "long_pct": round(100 * long / total, 1),
            "very_long_pct": round(100 * very_long / total, 1),
        },
        "transitions": {
            "short_after_long_pct": round(100 * short_after_long / transitions, 1),
            "long_after_short_pct": round(100 * long_after_short / transitions, 1),
            "same_band_pct": round(100 * same_band / transitions, 1),
        },
        "paratactic_pct": round(100 * paratactic / total, 1),
    }


def display_source_profile(profile):
    st.markdown("### Source Text Rhythm Profile")
    col1, col2, col3 = st.columns(3)
    col1.metric("Sentence CV", profile["cv"], help="Target for chapter rhythm")
    col2.metric("Mean length", f"{profile['mean_length']} words")
    col3.metric("Range", f"{profile['min_length']}–{profile['max_length']} words")
    
    col4, col5, col6 = st.columns(3)
    col4.metric("Short after long", f"{profile['transitions']['short_after_long_pct']}%",
                help="% of times a ≤5-word sentence follows a ≥25-word sentence")
    col5.metric("Same-band consecutive", f"{profile['transitions']['same_band_pct']}%",
                help="Lower = more variation")
    col6.metric("Paratactic (3+ 'and')", f"{profile['paratactic_pct']}%")
    
    d = profile["distribution"]
    st.caption(f"Distribution: ≤5w: {d['ultra_short_pct']}% | 6-12w: {d['short_pct']}% | "
               f"13-25w: {d['medium_pct']}% | 26-40w: {d['long_pct']}% | >40w: {d['very_long_pct']}%")


def build_rhythm_instructions(source_profile, chapter_cv):
    if source_profile is None:
        return ""
    
    instructions = []
    src_cv = source_profile["cv"]
    
    if chapter_cv < src_cv - 0.1:
        gap = round(src_cv - chapter_cv, 2)
        instructions.append(
            f"SENTENCE RHYTHM: The source authors have a sentence length CV of {src_cv}. "
            f"This chapter is at {chapter_cv} — {gap} points too uniform. To fix this:"
        )
        src_d = source_profile["distribution"]
        instructions.append(
            f"- Source distribution: {src_d['ultra_short_pct']}% ultra-short (≤5 words), "
            f"{src_d['very_long_pct']}% very long (>40 words). "
            f"The chapter needs MORE of both extremes."
        )
        src_t = source_profile["transitions"]
        if src_t["short_after_long_pct"] > 5:
            instructions.append(
                f"- In the source texts, {src_t['short_after_long_pct']}% of sentences ≤5 words follow sentences ≥25 words. "
                f"After a long sentence, follow it with something brutally short: 'She went.' / 'I knew.' / 'Small mercies.'"
            )
        if source_profile["paratactic_pct"] > 2:
            instructions.append(
                f"- Source texts use paratactic accumulation {source_profile['paratactic_pct']}% of the time. "
                f"Occasionally fuse two medium sentences into one long one using 'and...and...and' rhythm."
            )
        instructions.append(
            "- Do NOT just break long sentences into medium ones — that makes CV worse."
        )
    
    return "\n".join(instructions)


# ──────────────────────────────────────────────
# DETECTION SCORING ENGINE
# ──────────────────────────────────────────────

def score_chapter(text):
    """
    Scores a chapter against 12 detection metrics derived from
    reverse-engineering Originality.ai's Lite 1.0.2 detector.
    """
    words = text.split()
    word_count = len(words)
    if word_count == 0:
        return None
    
    kw = word_count / 1000
    
    sentence_endings = re.split(r'(?<=[.!?])\s+(?=[A-Z"\u201C])', text)
    sentences = [s.strip() for s in sentence_endings if s.strip()]
    sentence_lengths = [len(s.split()) for s in sentences]
    
    # ── 1. Em dash density ──
    em_dashes = text.count('\u2014') + text.count('--')
    em_dash_rate = em_dashes / kw if kw > 0 else 0
    
    # ── 2. "As though" / "as if" density ──
    as_though_count = len(re.findall(r'\bas though\b', text, re.I))
    as_if_count = len(re.findall(r'\bas if\b', text, re.I))
    as_though_rate = (as_though_count + as_if_count) / kw if kw > 0 else 0
    
    # ── 3. "The way he/she/they/I/men/people" density ──
    the_way_count = len(re.findall(r'\bthe way (?:he|she|they|I|it|men|women|people|a man|a woman|soldiers|hungry men|anyone|you|one)\b', text, re.I))
    the_way_rate = the_way_count / kw if kw > 0 else 0
    
    # ── 4. Negation-leading constructions ──
    negation_patterns = [
        r'\bIt was not\b', r'\bI did not\b', r'\bShe did not\b', r'\bHe did not\b',
        r'\bThat was not\b', r'\bThis was not\b', r'\bI had not\b',
        r'\bI was not\b', r'\bShe was not\b', r'\bHe was not\b',
        r'\bI could not\b', r'\bIt did not\b'
    ]
    negation_count = sum(len(re.findall(p, text)) for p in negation_patterns)
    negation_rate = negation_count / kw if kw > 0 else 0
    
    # ── 5. Period-to-comma ratio ──
    periods = text.count('.')
    commas = text.count(',')
    period_comma_ratio = periods / commas if commas > 0 else periods
    
    # ── 6. Dialogue density ──
    dialogue_matches = re.findall(r'[\u201C"][^"\u201D]*[\u201D"]', text)
    dialogue_words = sum(len(m.split()) for m in dialogue_matches)
    dialogue_pct = (dialogue_words / word_count * 100) if word_count > 0 else 0
    
    # ── 7. Sentence length variation (CV) ──
    if len(sentence_lengths) > 1:
        mean_len = sum(sentence_lengths) / len(sentence_lengths)
        variance = sum((x - mean_len) ** 2 for x in sentence_lengths) / len(sentence_lengths)
        std_dev = math.sqrt(variance)
        cv = std_dev / mean_len if mean_len > 0 else 0
    else:
        cv = 0
        mean_len = sentence_lengths[0] if sentence_lengths else 0
    
    # ── 8. Metacognitive verbs ──
    meta_verbs = re.findall(r'\b(?:I\s+)?(?:noted|filed|registered|understood|recognised|recognized|observed)\b', text, re.I)
    meta_rate = len(meta_verbs) / kw if kw > 0 else 0
    
    # ── 9. "The fact that" and "not X but Y" ──
    fact_that = len(re.findall(r'\bthe fact that\b', text, re.I))
    not_x_but_y = len(re.findall(r'\bnot [a-z]+ but [a-z]+\b', text, re.I))
    analytical_frames = fact_that + not_x_but_y
    analytical_rate = analytical_frames / kw if kw > 0 else 0
    
    # ── 10. "Of a man/woman/person who" characterization ──
    of_person_who = len(re.findall(r'\bof (?:a |someone |a man |a woman |a person |the kind of |the sort of )?(?:man|woman|person|someone|people) who\b', text, re.I))
    of_person_rate = of_person_who / kw if kw > 0 else 0
    
    # ── 11. Constructed similes and metaphors ──
    simile_patterns = [
        r'\bas though\b',
        r'\bas if\b',
        r'\bthe way (?:he|she|they|I|it|men|women|people|a man|a woman|soldiers|hungry|anyone|you|one)\b',
        r'\blike a [a-z]+ (?:that|who|which)\b',
        r'\blike a [a-z]+ [a-z]+ing\b',
        r'\blike an? [a-z]+\b',                     # "like an invalid", "like a prayer"
        r'\bthe kind of [a-z]+ (?:that|who|which|you)\b',
        r'\bthe sort of [a-z]+ (?:that|who|which|you)\b',
        r'\b[a-z]+ed the way [a-z]+\b',
        r'\bas a [a-z]+ (?:does|would|might|could|who)\b',
        r'\bthe particular [a-z]+ (?:of|that|I|she|he|it)\b',
        r'\bwith the [a-z]+ [a-z]+ of (?:a|someone|an)\b',    # "with the calm authority of a"
        r'\bwith the [a-z]+ of a\b',                            # "with the patience of a"
        r'\bhad the [a-z]+ of a\b',
        r'\bin the manner of\b',
        r'\bwith the air of\b',
        r'\bthe [a-z]+ of survival\b',              # "the arithmetic of survival"
        r'\bthe [a-z]+ of (?:male |female )?(?:pride|fear|grief|joy|despair|hope)\b',  # "the momentum of male pride"
    ]
    simile_count = 0
    simile_matches_all = []
    for pat in simile_patterns:
        found = re.findall(pat, text, re.I)
        simile_count += len(found)
        simile_matches_all.extend(found)
    simile_rate = simile_count / kw if kw > 0 else 0
    
    # ── 12. Editorial commentary ──
    # Sentences where the narrator explains what just happened or 
    # draws a moral/conclusion the reader can draw themselves.
    editorial_patterns = [
        r'^This was (?:what|how|the)\b',
        r'^That was (?:what|how|the)\b',
        r'\bwhich meant (?:that |neither |both |every |we |she |he |I )\b',
        r'\band that made (?:all )?the difference\b',
        r'\bthe arithmetic of\b',
        r'\bthe mathematics of\b',
        r'\bevery (?:small )?(?:victory|compromise|expense|decision|loss) (?:counted|mattered|had consequences)\b',
        r'\b(?:neither|both) of us (?:was|were|could|had|knew)\b',
        r'\bThis was the (?:foundation|arithmetic|mathematics|logic|rhythm|pattern|way)\b',
    ]
    editorial_count = 0
    for sent in sentences:
        for pat in editorial_patterns:
            if re.search(pat, sent, re.I):
                editorial_count += 1
                break
    editorial_rate = editorial_count / kw if kw > 0 else 0
    
    # ── Scoring ──
    metrics = {}
    
    def rate(name, value, green_thresh, yellow_thresh, unit, invert=False):
        if invert:
            if value >= green_thresh:
                level = "GREEN"
            elif value >= yellow_thresh:
                level = "YELLOW"
            else:
                level = "RED"
        else:
            if value <= green_thresh:
                level = "GREEN"
            elif value <= yellow_thresh:
                level = "YELLOW"
            else:
                level = "RED"
        metrics[name] = {"value": round(value, 3), "level": level, "unit": unit}
    
    rate("Em dash density", em_dash_rate, 1.0, 2.5, "/1000w")
    rate("'As though/as if' density", as_though_rate, 0.2, 0.6, "/1000w")
    rate("'The way he/she' density", the_way_rate, 0.5, 1.0, "/1000w")
    rate("Negation-leading density", negation_rate, 2.0, 4.0, "/1000w")
    rate("Period-to-comma ratio", period_comma_ratio, 1.8, 1.5, "ratio", invert=True)
    rate("Dialogue density", dialogue_pct, 15.0, 8.0, "% of words", invert=True)
    rate("Sentence length CV", cv, 1.2, 1.0, "coefficient", invert=True)
    rate("Metacognitive verb density", meta_rate, 0.3, 1.0, "/1000w")
    rate("Analytical frames", analytical_rate, 0.0, 0.2, "/1000w")
    rate("'Of a person who' density", of_person_rate, 0.0, 0.3, "/1000w")
    rate("Constructed simile density", simile_rate, 1.0, 3.0, "/1000w")
    rate("Editorial commentary", editorial_rate, 0.0, 0.3, "/1000w")
    
    # ── Flag specific high-risk passages ──
    flagged = []
    
    for i, sent in enumerate(sentences):
        risks = []
        
        if '\u2014' in sent or '--' in sent:
            risks.append("em_dash")
        
        if re.search(r'\bas though\b|\bas if\b', sent, re.I):
            risks.append("as_though")
        
        if re.search(r'\bthe way (?:he|she|they|I|it|men|women|people|a man|a woman|soldiers|hungry|anyone|you|one)\b', sent, re.I):
            risks.append("the_way")
        
        if re.search(r'\bof (?:a |someone |a man |a woman |a person )?(?:man|woman|person|someone|people) who\b', sent, re.I):
            risks.append("of_person_who")
        
        if re.search(r'\b(?:I\s+)?(?:noted|filed|registered|understood|recognised|recognized)\b', sent, re.I):
            risks.append("metacognitive")
        
        if re.search(r'\bthe fact that\b', sent, re.I):
            risks.append("fact_that")
        
        if re.search(r'^(?:It|That|This|I|She|He) (?:was|did|had|could) not\b', sent):
            risks.append("negation_leading")
        
        comma_count = sent.count(',')
        sent_words = len(sent.split())
        if comma_count >= 4 and sent_words >= 35:
            risks.append("long_compound")
        
        if ('\u2014' in sent or '--' in sent) and sent_words > 25:
            if re.search(r'\bas though\b|\bthe way\b|\bwhich (?:was|meant|told)\b', sent, re.I):
                risks.append("obs_interp_coupling")
        
        # Constructed simile at sentence level
        simile_sentence_patterns = [
            r'\blike a [a-z]+ (?:that|who|which)\b',
            r'\blike a [a-z]+ [a-z]+ing\b',
            r'\blike an? [a-z]+\b',
            r'\bthe kind of [a-z]+ (?:that|who|which|you)\b',
            r'\bthe sort of [a-z]+ (?:that|who|which|you)\b',
            r'\b[a-z]+ed the way [a-z]+\b',
            r'\bas a [a-z]+ (?:does|would|might|could|who)\b',
            r'\bwith the [a-z]+ [a-z]+ of (?:a|someone|an)\b',
            r'\bwith the [a-z]+ of a\b',
            r'\bhad the [a-z]+ of a\b',
            r'\bin the manner of\b',
            r'\bwith the air of\b',
            r'\bthe particular [a-z]+ (?:of|that|I|she|he|it)\b',
            r'\bthe [a-z]+ of survival\b',
            r'\bthe [a-z]+ of (?:male |female )?(?:pride|fear|grief|joy|despair|hope)\b',
        ]
        for sp in simile_sentence_patterns:
            if re.search(sp, sent, re.I):
                risks.append("constructed_simile")
                break
        
        # Editorial commentary at sentence level
        for pat in editorial_patterns:
            if re.search(pat, sent, re.I):
                risks.append("editorial")
                break
        
        if risks:
            flagged.append({
                "index": i,
                "sentence": sent,
                "risks": risks,
                "risk_count": len(risks)
            })
    
    # ── Overall risk score ──
    red_count = sum(1 for m in metrics.values() if m["level"] == "RED")
    yellow_count = sum(1 for m in metrics.values() if m["level"] == "YELLOW")
    green_count = sum(1 for m in metrics.values() if m["level"] == "GREEN")
    
    if red_count >= 4:
        overall = "HIGH RISK"
    elif red_count >= 2 or (red_count >= 1 and yellow_count >= 3):
        overall = "MODERATE RISK"
    elif red_count == 0 and yellow_count <= 2:
        overall = "LOW RISK"
    else:
        overall = "MODERATE RISK"
    
    return {
        "word_count": word_count,
        "sentence_count": len(sentences),
        "metrics": metrics,
        "flagged": sorted(flagged, key=lambda x: -x["risk_count"]),
        "overall": overall,
        "red_count": red_count,
        "yellow_count": yellow_count,
        "green_count": green_count,
        "summary": {
            "em_dashes": em_dashes,
            "as_though_total": as_though_count + as_if_count,
            "the_way_total": the_way_count,
            "negation_total": negation_count,
            "simile_total": simile_count,
            "editorial_total": editorial_count,
            "dialogue_word_pct": round(dialogue_pct, 1),
            "mean_sentence_length": round(mean_len, 1),
            "flagged_sentences": len(flagged),
            "total_sentences": len(sentences)
        }
    }


# ──────────────────────────────────────────────
# THREE-PASS ARCHITECTURE: WRITE → CUT → FILL
# ──────────────────────────────────────────────

def build_cutting_prompt(chapter_text, score_result):
    """
    PASS 2: THE BLADE.
    
    Identifies flagged passages and tells the model to DELETE them.
    No rewriting. No replacement imagery. Just removal and gap closure.
    The model outputs the complete chapter with flagged material excised.
    """
    flagged = score_result["flagged"]
    metrics = score_result["metrics"]
    
    top_flagged = flagged[:25]
    
    # Build the list of passages to cut
    passage_list = ""
    for i, item in enumerate(top_flagged):
        risk_labels = ", ".join(item["risks"])
        passage_list += f'\n\nCUT {i+1} [{risk_labels}]:\n"{item["sentence"][:300]}"'
    
    metric_warnings = ""
    for name, data in metrics.items():
        if data["level"] == "RED":
            metric_warnings += f"\n- {name}: {data['value']} {data['unit']} — RED"
        elif data["level"] == "YELLOW":
            metric_warnings += f"\n- {name}: {data['value']} {data['unit']} — YELLOW"
    
    prompt = f"""You are performing a CUTTING PASS on a chapter of fiction. Your job is to DELETE material that triggers AI detection. You are a blade, not a writer.

RULES — READ CAREFULLY:

1. DEFAULT ACTION IS DELETE. When you encounter a flagged passage, your first instinct must be to remove it entirely. Not rewrite it. Not rephrase it. Delete it.

2. CLOSE THE GAP. After deleting a passage, connect the sentence before it to the sentence after it. If the connection reads naturally, you are done. If the connection needs a bridge, write ONE short plain sentence (under 10 words) to smooth the transition. Nothing more.

3. NEVER GENERATE REPLACEMENT IMAGERY. Do not replace a deleted simile with a new simile. Do not replace a deleted metaphor with a new metaphor. Do not replace editorial commentary with different editorial commentary. The replacement for a constructed image is NOTHING.

4. SPECIFIC CONSTRUCTIONS TO DELETE ON SIGHT:
   - "with the [adjective] [noun] of someone/a man/a woman who..." → DELETE the entire phrase. Keep only the action it was attached to.
   - "the kind/sort of [noun] that/who..." → DELETE. Replace with the plain noun if needed.
   - "like a [noun] that/who..." → DELETE the comparison entirely.
   - "as though [interpretation]" → DELETE everything after "as though."
   - "the particular [noun] of/that..." → DELETE. Use a plain adjective or nothing.
   - "the [abstract noun] of [abstract noun]" (e.g., "the arithmetic of survival", "the momentum of male pride") → DELETE the entire phrase.
   - "with the air of..." → DELETE.
   - Any sentence beginning with "This was what..." or "This was how..." or "That was the..." where the narrator explains what just happened → DELETE THE ENTIRE SENTENCE.
   - Any sentence containing "every [noun] mattered/counted/had consequences" → DELETE THE ENTIRE SENTENCE.
   - Any sentence containing "neither of us" + editorial observation → DELETE THE ENTIRE SENTENCE.
   - Any sentence that interprets a scene the reader has already witnessed → DELETE IT.

5. PRESERVE EVERYTHING ELSE. Do not touch dialogue. Do not touch physical action. Do not touch scene transitions. Do not touch plot beats. Only cut what is flagged or matches the patterns above.

6. THE CHAPTER WILL BE SHORTER AFTER THIS PASS. That is correct. Do not try to maintain word count. The next pass will add new material. Your job is only to cut.

7. OUTPUT THE COMPLETE CHAPTER with your cuts applied. Every sentence that is NOT flagged or matched to the patterns above must appear in your output exactly as written.

CURRENT METRICS:
{metric_warnings}

PASSAGES TO CUT:
{passage_list}

IMPORTANT: Also scan the ENTIRE chapter for any instances of the patterns in rule 4 that were not flagged above. The flagging system catches most but not all. If you see "with the calm authority of someone who had done this a hundred times" anywhere in the text, cut it even if it was not in the flagged list.

<chapter>
{chapter_text}
</chapter>

Output the complete chapter with all cuts applied. Do not add any commentary before or after the chapter text."""
    
    return prompt


def build_fill_prompt(cut_chapter, original_word_count, prompt_text, outline_text, source_text, char_text, source_profile=None):
    """
    PASS 3: THE FILL.
    
    Takes the cut chapter (shorter than original), the original writing prompt,
    outline, source texts, and character profiles. Generates new material
    to restore word count, anchored to the outline's microbeats and the
    source texts' voice — not to the model's default literary instincts.
    """
    cut_word_count = len(cut_chapter.split())
    deficit = original_word_count - cut_word_count
    
    # Build rhythm instructions if we have source profile
    chapter_cv = 0
    sentence_endings = re.split(r'(?<=[.!?])\s+(?=[A-Z"\u201C])', cut_chapter)
    sents = [s.strip() for s in sentence_endings if s.strip()]
    if len(sents) > 1:
        lengths = [len(s.split()) for s in sents]
        m = sum(lengths) / len(lengths)
        v = sum((x - m) ** 2 for x in lengths) / len(lengths)
        chapter_cv = round(math.sqrt(v) / m, 3) if m > 0 else 0
    
    rhythm_block = build_rhythm_instructions(source_profile, chapter_cv) if source_profile else ""
    
    prompt = f"""You are performing a FILL PASS on a chapter of fiction. The chapter has been through a cutting pass that removed AI-detectable constructions (similes, metaphors, editorial commentary, interpretive glosses). It is now {cut_word_count:,} words. The original was {original_word_count:,} words. You need to add approximately {deficit} words of new material.

YOUR ANCHOR DOCUMENTS:

The chapter outline below specifies every microbeat in the chapter — what happens, what the characters say, what the narrator notices, and what the narrator does NOT do. The source texts below are your voice model. The character profiles below tell you who these people are.

RULES — READ CAREFULLY:

1. IDENTIFY THE THIN SPOTS. Read the cut chapter against the outline. Where did the cuts leave two scene beats too close together? Where did a transition disappear? Where did a physical detail get lost? Where is a SCENE-weight microbeat now underserved? Those are your insertion points.

2. WRITE NEW MATERIAL FOR THOSE SPOTS. Not replacement material — new material. You have never seen the deleted text. You are writing fresh sentences that belong in these gaps, per the outline's specifications for each microbeat.

3. ANCHOR TO THE SOURCE TEXTS. Before writing each new sentence, silently identify which source passage's rhythm, construction, and punctuation you are channeling. Do not cite it. Do not reproduce it. Let it guide the sentence's shape, then write original prose in the narrator's voice. The source passage is the blueprint. The words are yours.

4. MATCH THE SURROUNDING VOICE. For each insertion point, read the sentence immediately before and immediately after the gap. Your new material must sound like it was written in the same breath — same register, same rhythm, same level of plainness or compression. Do not shift gears.

5. BANNED CONSTRUCTIONS — DO NOT USE ANY OF THESE:
   - Em dashes (use periods and new sentences instead)
   - "as though" / "as if" + interpretation
   - "the way he/she/they" + interpretation
   - "with the [adjective] [noun] of someone/a man/a woman who"
   - "the kind/sort of [noun] that/who"
   - "like a [noun] that/who" (constructed similes)
   - "the particular [noun] of"
   - "the [abstract noun] of [abstract noun]" (metaphorical equations)
   - "with the air of"
   - Sentences beginning "This was what/how" (editorial summaries)
   - Sentences that explain, interpret, or editorialize about what just happened
   - Metacognitive verbs: noted, registered, filed, observed, understood, recognised
   - "of a man/woman/person who" characterization clauses
   - "not X but Y" analytical framing
   - "the fact that"
   - Any sentence that tells the reader what to think about an action they just witnessed

6. WHAT TO WRITE INSTEAD:
   - Physical action. What does the character do with her hands, her feet, her eyes?
   - Dialogue. The outline specifies weapons and resolutions for each exchange. Write more of the exchange.
   - Specific sensory detail. Not atmosphere — specific objects. A nail, a stain, a sound, a smell.
   - Dinah's opinions. She has opinions about everything. They are short, direct, and ungenerous. They do not philosophise.
   - Plain declarative sentences. "The bolt stuck." "He left." "I counted twice."
   - Vary sentence lengths aggressively. Follow a long sentence with a 2-5 word sentence. Follow a short sentence with a long accumulating clause using "and...and...and" rhythm.

{rhythm_block}

7. SEAMLESS MELDING. The reader must not be able to tell where the old material ends and the new material begins. Read the sentence before, write your insertion, read the sentence after. If the join is visible, rewrite the join until it disappears.

8. OUTPUT THE COMPLETE CHAPTER with all new material woven in. Every original sentence that survived the cut must appear exactly as written. Your new material sits between them.

<source_texts>
{source_text if source_text else "(No source texts provided)"}
</source_texts>

<character_profiles>
{char_text if char_text else "(No character profiles provided)"}
</character_profiles>

<chapter_outline>
{outline_text}
</chapter_outline>

<writing_prompt>
{prompt_text}
</writing_prompt>

<cut_chapter>
{cut_chapter}
</cut_chapter>

Output the complete chapter with new material seamlessly woven in. Do not add any commentary before or after the chapter text."""
    
    return prompt


# Legacy revision prompt for manual revise button
def build_revision_prompt(chapter_text, score_result, source_profile=None):
    """Backward-compatible revision prompt for the manual 'Revise This Text' button."""
    return build_cutting_prompt(chapter_text, score_result)


def generate_report(score_result, chapter_text, label, pass_num, source_profile=None):
    """Generates a Word document report for a scoring pass."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    title = doc.add_heading(f'Detection Score Report — {label}', level=1)
    
    summary = score_result["summary"]
    overall = score_result["overall"]
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    
    p = doc.add_paragraph()
    p.add_run(f"Generated: {timestamp}\n").bold = False
    p.add_run(f"Words: {score_result['word_count']:,} | "
              f"Sentences: {summary['total_sentences']} | "
              f"Mean sentence length: {summary['mean_sentence_length']} words\n")
    p.add_run(f"Dialogue: {summary['dialogue_word_pct']}% of words | "
              f"Flagged sentences: {summary['flagged_sentences']} of {summary['total_sentences']} "
              f"({round(100*summary['flagged_sentences']/max(summary['total_sentences'],1))}%)\n")
    
    p2 = doc.add_paragraph()
    risk_run = p2.add_run(f"OVERALL: {overall}")
    risk_run.bold = True
    risk_run.font.size = Pt(14)
    if overall == "HIGH RISK":
        risk_run.font.color.rgb = RGBColor(200, 0, 0)
    elif overall == "MODERATE RISK":
        risk_run.font.color.rgb = RGBColor(200, 150, 0)
    else:
        risk_run.font.color.rgb = RGBColor(0, 150, 0)
    
    p2.add_run(f"  ({score_result['red_count']} RED, {score_result['yellow_count']} YELLOW, {score_result['green_count']} GREEN)")
    
    doc.add_heading('Metrics', level=2)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Value'
    hdr[2].text = 'Unit'
    hdr[3].text = 'Status'
    
    for name, data in score_result["metrics"].items():
        row = table.add_row().cells
        row[0].text = name
        row[1].text = str(data["value"])
        row[2].text = data["unit"]
        row[3].text = data["level"]
        for paragraph in row[3].paragraphs:
            for run in paragraph.runs:
                if data["level"] == "RED":
                    run.font.color.rgb = RGBColor(200, 0, 0)
                    run.bold = True
                elif data["level"] == "YELLOW":
                    run.font.color.rgb = RGBColor(200, 150, 0)
                    run.bold = True
                else:
                    run.font.color.rgb = RGBColor(0, 130, 0)
    
    doc.add_paragraph(
        f"Em dashes: {summary['em_dashes']} | "
        f"'As though/if': {summary['as_though_total']} | "
        f"'The way': {summary['the_way_total']} | "
        f"Negation-leading: {summary['negation_total']} | "
        f"Similes: {summary['simile_total']} | "
        f"Editorial: {summary['editorial_total']}"
    )
    
    if source_profile:
        doc.add_heading('Sentence Rhythm — Source Comparison', level=2)
        ch_cv = score_result["metrics"].get("Sentence length CV", {}).get("value", 0)
        src_cv = source_profile["cv"]
        src_d = source_profile["distribution"]
        src_t = source_profile["transitions"]
        
        rtable = doc.add_table(rows=1, cols=3)
        rtable.style = 'Light Grid Accent 1'
        rhdr = rtable.rows[0].cells
        rhdr[0].text = 'Measure'
        rhdr[1].text = 'Source'
        rhdr[2].text = 'Chapter'
        
        rhythm_rows = [
            ("Sentence CV", str(src_cv), str(ch_cv)),
            ("Mean sentence length", f"{source_profile['mean_length']}w", f"{summary['mean_sentence_length']}w"),
            ("Range (min-max)", f"{source_profile['min_length']}-{source_profile['max_length']}w", "-"),
            ("Ultra-short (<=5w)", f"{src_d['ultra_short_pct']}%", "-"),
            ("Short (6-12w)", f"{src_d['short_pct']}%", "-"),
            ("Medium (13-25w)", f"{src_d['medium_pct']}%", "-"),
            ("Long (26-40w)", f"{src_d['long_pct']}%", "-"),
            ("Very long (>40w)", f"{src_d['very_long_pct']}%", "-"),
            ("Short-after-long", f"{src_t['short_after_long_pct']}%", "-"),
            ("Same-band consecutive", f"{src_t['same_band_pct']}%", "-"),
            ("Paratactic (3+ 'and')", f"{source_profile['paratactic_pct']}%", "-"),
        ]
        for measure, src_val, ch_val in rhythm_rows:
            row = rtable.add_row().cells
            row[0].text = measure
            row[1].text = src_val
            row[2].text = ch_val
        
        gap = round(src_cv - ch_cv, 3)
        if gap > 0.1:
            p_rhythm = doc.add_paragraph()
            r = p_rhythm.add_run(f"Rhythm gap: {gap}")
            r.font.color.rgb = RGBColor(200, 150, 0)
            r.bold = True
    
    flagged = score_result["flagged"]
    if flagged:
        doc.add_heading(f'Flagged Passages ({len(flagged)})', level=2)
        for item in flagged[:25]:
            p = doc.add_paragraph()
            tag_run = p.add_run(f"[{item['risk_count']} flags: {', '.join(item['risks'])}]")
            tag_run.bold = True
            tag_run.font.size = Pt(9)
            tag_run.font.color.rgb = RGBColor(180, 0, 0)
            p2 = doc.add_paragraph()
            text_run = p2.add_run(item['sentence'][:400])
            text_run.font.size = Pt(10)
            text_run.italic = True
            doc.add_paragraph()
    
    doc.add_heading('Full Chapter Text', level=2)
    for para_text in chapter_text.split("\n"):
        if para_text.strip():
            doc.add_paragraph(para_text)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def display_scorecard(score_result, source_profile=None):
    st.markdown("### Detection Risk Scorecard")
    
    overall = score_result["overall"]
    if overall == "HIGH RISK":
        st.error(f"Overall: {overall} — {score_result['red_count']} red, {score_result['yellow_count']} yellow, {score_result['green_count']} green")
    elif overall == "MODERATE RISK":
        st.warning(f"Overall: {overall} — {score_result['red_count']} red, {score_result['yellow_count']} yellow, {score_result['green_count']} green")
    else:
        st.success(f"Overall: {overall} — {score_result['red_count']} red, {score_result['yellow_count']} yellow, {score_result['green_count']} green")
    
    col1, col2, col3, col4 = st.columns(4)
    metric_items = list(score_result["metrics"].items())
    
    for i, (name, data) in enumerate(metric_items):
        target_col = [col1, col2, col3, col4][i % 4]
        icon = {"GREEN": "\u2705", "YELLOW": "\u26A0\uFE0F", "RED": "\u274C"}[data["level"]]
        target_col.metric(
            label=f"{icon} {name}",
            value=f"{data['value']} {data['unit']}",
            delta=data["level"],
            delta_color="normal" if data["level"] == "GREEN" else ("off" if data["level"] == "YELLOW" else "inverse")
        )
    
    s = score_result["summary"]
    st.caption(f"{s['total_sentences']} sentences | Mean length: {s['mean_sentence_length']} words | "
               f"{s['flagged_sentences']} flagged ({round(100*s['flagged_sentences']/max(s['total_sentences'],1))}%) | "
               f"Dialogue: {s['dialogue_word_pct']}% | Similes: {s['simile_total']} | Editorial: {s['editorial_total']}")
    
    if source_profile:
        ch_cv = score_result["metrics"].get("Sentence length CV", {}).get("value", 0)
        src_cv = source_profile["cv"]
        src_d = source_profile["distribution"]
        gap = round(src_cv - ch_cv, 3)
        
        if gap > 0.1:
            st.warning(
                f"**Rhythm gap:** Chapter CV {ch_cv} vs Source CV {src_cv} (gap: {gap}). "
                f"Source has {src_d['ultra_short_pct']}% ultra-short (<=5w) and "
                f"{src_d['very_long_pct']}% very long (>40w)."
            )
        elif gap > 0:
            st.info(f"Rhythm close to source: Chapter CV {ch_cv} vs Source CV {src_cv}.")
        else:
            st.success(f"Rhythm matches or exceeds source: Chapter CV {ch_cv} vs Source CV {src_cv}.")
    
    flagged = score_result["flagged"]
    if flagged:
        with st.expander(f"Flagged Passages ({len(flagged)} sentences)", expanded=False):
            for item in flagged[:20]:
                risk_tags = " ".join([f"`{r}`" for r in item["risks"]])
                st.markdown(f"**[{item['risk_count']} flags]** {risk_tags}")
                st.text(item["sentence"][:200])
                st.markdown("---")


# ──────────────────────────────────────────────
# SIDEBAR
# ──────────────────────────────────────────────

with st.sidebar:
    st.header("Settings")
    api_key = st.text_input("Anthropic API Key", type="password")
    
    model_choice = st.selectbox("Writing Model", [
        "Sonnet",
        "Sonnet Extended Thinking",
        "Opus",
        "Haiku"
    ])
    
    model_map = {
        "Sonnet": "claude-sonnet-4-20250514",
        "Sonnet Extended Thinking": "claude-sonnet-4-20250514",
        "Opus": "claude-opus-4-20250514",
        "Haiku": "claude-haiku-4-5-20251001"
    }
    
    model_id = model_map[model_choice]
    
    temperature = st.slider("Temperature", 0.0, 1.0, 1.0, 0.1)
    max_tokens = st.number_input("Max Tokens", min_value=1000, max_value=128000, value=16000, step=1000)
    
    if model_choice == "Sonnet Extended Thinking":
        thinking_budget = st.number_input("Thinking Budget (tokens)", min_value=1000, max_value=50000, value=10000, step=1000)
        st.caption("Extended thinking requires temperature = 1.0.")
    
    st.markdown("---")
    st.header("Three-Pass Pipeline")
    st.caption("WRITE → CUT → FILL")
    auto_pipeline = st.checkbox("Run full pipeline (Write + Cut + Fill)", value=False)
    run_cut_pass = st.checkbox("Run cutting pass", value=True, help="Deletes AI-detectable constructions")
    run_fill_pass = st.checkbox("Run fill pass", value=True, help="Adds new material per outline to restore word count")
    
    st.markdown("---")
    st.header("Cut/Fill Model")
    revision_model_choice = st.selectbox("Model for Cut and Fill passes", [
        "Same as writing model",
        "Sonnet",
        "Sonnet Extended Thinking",
        "Haiku"
    ])
    
    if revision_model_choice == "Sonnet Extended Thinking":
        rev_thinking_budget = st.number_input("Cut/Fill Thinking Budget", min_value=1000, max_value=50000, value=10000, step=1000, key="rev_think")


# ──────────────────────────────────────────────
# MAIN AREA
# ──────────────────────────────────────────────

st.subheader("Source Documents")
col1, col2, col3 = st.columns(3)
with col1:
    source_file = st.file_uploader("Source Texts", type=["txt", "docx"])
with col2:
    char_file = st.file_uploader("Character Profiles", type=["txt", "docx"])
with col3:
    outline_file = st.file_uploader("Chapter Outline", type=["txt", "docx"])

def read_uploaded(f):
    if f is None:
        return ""
    if f.name.endswith(".txt"):
        return f.read().decode("utf-8")
    elif f.name.endswith(".docx"):
        doc = Document(io.BytesIO(f.read()))
        return "\n".join([p.text for p in doc.paragraphs])
    return ""

prompt_default = """Using the source texts, character profiles, and chapter outline provided, write the chapter in one continuous pass from first sentence to last. Do not draft short and expand.

Before writing each sentence, silently identify which source passage's rhythm, construction, and punctuation you are channeling. Do not cite it. Do not reproduce it. Let it guide the sentence's shape, then write original prose in the narrator's voice. The source passage is the blueprint. The words are yours.

Do not construct sentences that explain, interpret, or editorialize. Trust the image. Trust the action. Trust the reader."""

prompt = st.text_area("Writing Prompt", value=prompt_default, height=150)

# ── Analyze source texts when uploaded ──
if source_file is not None:
    source_text_for_analysis = read_uploaded(source_file)
    source_file.seek(0)
    if source_text_for_analysis.strip():
        profile = analyze_source_rhythm(source_text_for_analysis)
        if profile:
            st.session_state.source_profile = profile
            with st.expander("Source Rhythm Profile", expanded=False):
                display_source_profile(profile)
        else:
            st.caption("Source text too short for rhythm analysis (need 10+ sentences).")
else:
    st.session_state.source_profile = None

# ──────────────────────────────────────────────
# SESSION STATE
# ──────────────────────────────────────────────

if "chapter_text" not in st.session_state:
    st.session_state.chapter_text = None
if "score_result" not in st.session_state:
    st.session_state.score_result = None
if "revision_history" not in st.session_state:
    st.session_state.revision_history = []
if "current_pass" not in st.session_state:
    st.session_state.current_pass = 0
if "reports" not in st.session_state:
    st.session_state.reports = []
if "source_profile" not in st.session_state:
    st.session_state.source_profile = None


def call_api(client, message_text, is_revision=False):
    """Makes an API call with the current settings."""
    
    # Determine model for this call
    if is_revision and revision_model_choice != "Same as writing model":
        rev_model = model_map.get(revision_model_choice, model_id)
        use_extended = (revision_model_choice == "Sonnet Extended Thinking")
        
        if use_extended:
            budget = rev_thinking_budget if 'rev_thinking_budget' in dir() else 10000
            response = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=max_tokens,
                temperature=1.0,
                thinking={"type": "enabled", "budget_tokens": budget},
                messages=[{"role": "user", "content": message_text}]
            )
        else:
            response = client.messages.create(
                model=rev_model,
                max_tokens=max_tokens,
                temperature=temperature,
                messages=[{"role": "user", "content": message_text}]
            )
    elif model_choice == "Sonnet Extended Thinking" and not is_revision:
        response = client.messages.create(
            model=model_id,
            max_tokens=max_tokens,
            temperature=1.0,
            thinking={"type": "enabled", "budget_tokens": thinking_budget},
            messages=[{"role": "user", "content": message_text}]
        )
    else:
        response = client.messages.create(
            model=model_id,
            max_tokens=max_tokens,
            temperature=temperature,
            messages=[{"role": "user", "content": message_text}]
        )
    
    chapter_text = ""
    thinking_text = ""
    for block in response.content:
        if block.type == "thinking":
            thinking_text = block.thinking
        elif block.type == "text":
            chapter_text += block.text
    
    return chapter_text, thinking_text


def make_docx(text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(12)
    for para_text in text.split("\n"):
        doc.add_paragraph(para_text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ──────────────────────────────────────────────
# WRITE BUTTON — THREE-PASS PIPELINE
# ──────────────────────────────────────────────

if st.button("Write Chapter", type="primary"):
    if not api_key:
        st.error("Enter your API key in the sidebar.")
    elif not outline_file:
        st.error("Upload at least a chapter outline.")
    else:
        source_text = read_uploaded(source_file)
        char_text = read_uploaded(char_file)
        outline_text = read_uploaded(outline_file)
        
        parts = []
        if source_text:
            parts.append(f"<source_texts>\n{source_text}\n</source_texts>")
        if char_text:
            parts.append(f"<character_profiles>\n{char_text}\n</character_profiles>")
        parts.append(f"<chapter_outline>\n{outline_text}\n</chapter_outline>")
        parts.append(prompt)
        
        user_message = "\n\n".join(parts)
        
        char_count = len(user_message)
        token_est = char_count // 4
        st.info(f"Input: ~{char_count:,} characters (~{token_est:,} tokens)")
        
        with st.spinner("PASS 1: Writing chapter..."):
            try:
                client = anthropic.Anthropic(api_key=api_key)
                chapter_text, thinking_text = call_api(client, user_message)
                
                if not chapter_text.strip():
                    st.error("Model returned empty response.")
                else:
                    if thinking_text:
                        with st.expander("Model's thinking process"):
                            st.text(thinking_text[:3000])
                    
                    original_word_count = len(chapter_text.split())
                    st.success(f"PASS 1 complete — {original_word_count:,} words")
                    
                    # Store
                    st.session_state.chapter_text = chapter_text
                    st.session_state.revision_history = [{"pass": 0, "text": chapter_text, "label": "Original (Write)"}]
                    st.session_state.current_pass = 0
                    st.session_state.reports = []
                    
                    # Score
                    score_result = score_chapter(chapter_text)
                    st.session_state.score_result = score_result
                    
                    # Report
                    report_buf = generate_report(score_result, chapter_text, "Original (Write)", 0, st.session_state.source_profile)
                    st.session_state.reports.append({"label": "Original (Write)", "buffer": report_buf})
                    
                    # Display
                    st.markdown("---")
                    st.markdown("### PASS 1: Original")
                    display_scorecard(score_result, st.session_state.source_profile)
                    
                    with st.expander("Original Chapter Text", expanded=False):
                        st.text(chapter_text)
                    
                    dcol1, dcol2 = st.columns(2)
                    with dcol1:
                        buffer = make_docx(chapter_text)
                        st.download_button(
                            label="Download Original (.docx)",
                            data=buffer,
                            file_name="chapter_original.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    with dcol2:
                        st.download_button(
                            label="Download Original Report (.docx)",
                            data=st.session_state.reports[-1]["buffer"],
                            file_name="report_original.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="report_dl_orig"
                        )
                    
                    current_text = chapter_text
                    current_score = score_result
                    
                    # ── PASS 2: CUT ──
                    if auto_pipeline and run_cut_pass and current_score["overall"] != "LOW RISK":
                        st.markdown("---")
                        st.markdown("### PASS 2: Cut")
                        
                        if not current_score["flagged"]:
                            st.info("No flagged passages. Skipping cut pass.")
                        else:
                            cutting_prompt = build_cutting_prompt(current_text, current_score)
                            
                            with st.spinner("PASS 2: Cutting AI-detectable constructions..."):
                                cut_text, cut_thinking = call_api(client, cutting_prompt, is_revision=True)
                            
                            if not cut_text.strip():
                                st.warning("Cut pass returned empty. Skipping.")
                            else:
                                if cut_thinking:
                                    with st.expander("Cut pass thinking"):
                                        st.text(cut_thinking[:3000])
                                
                                cut_word_count = len(cut_text.split())
                                deficit = original_word_count - cut_word_count
                                st.info(f"Cut pass: {original_word_count:,} → {cut_word_count:,} words ({deficit:,} words removed)")
                                
                                cut_score = score_chapter(cut_text)
                                
                                # Report
                                cut_report = generate_report(cut_score, cut_text, "After Cut", 1, st.session_state.source_profile)
                                st.session_state.reports.append({"label": "After Cut", "buffer": cut_report})
                                st.session_state.revision_history.append({"pass": 1, "text": cut_text, "label": "After Cut"})
                                
                                # Comparison
                                rcol1, rcol2, rcol3 = st.columns(3)
                                rcol1.metric("Red metrics", cut_score["red_count"], 
                                           delta=f"{cut_score['red_count'] - current_score['red_count']}", delta_color="inverse")
                                rcol2.metric("Flagged sentences", len(cut_score["flagged"]), 
                                           delta=f"{len(cut_score['flagged']) - len(current_score['flagged'])}", delta_color="inverse")
                                rcol3.metric("Words removed", f"{deficit:,}")
                                
                                display_scorecard(cut_score, st.session_state.source_profile)
                                
                                with st.expander("Cut Chapter Text", expanded=False):
                                    st.text(cut_text)
                                
                                pcol1, pcol2 = st.columns(2)
                                with pcol1:
                                    cut_buf = make_docx(cut_text)
                                    st.download_button(
                                        label="Download Cut Version (.docx)",
                                        data=cut_buf,
                                        file_name="chapter_after_cut.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="ch_dl_cut"
                                    )
                                with pcol2:
                                    st.download_button(
                                        label="Download Cut Report (.docx)",
                                        data=st.session_state.reports[-1]["buffer"],
                                        file_name="report_after_cut.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="report_dl_cut"
                                    )
                                
                                current_text = cut_text
                                current_score = cut_score
                                
                                # ── PASS 3: FILL ──
                                if run_fill_pass and deficit > 100:
                                    st.markdown("---")
                                    st.markdown("### PASS 3: Fill")
                                    
                                    fill_prompt = build_fill_prompt(
                                        current_text,
                                        original_word_count,
                                        prompt,
                                        outline_text,
                                        source_text,
                                        char_text,
                                        st.session_state.source_profile
                                    )
                                    
                                    with st.spinner(f"PASS 3: Filling {deficit:,} words of new material..."):
                                        fill_text, fill_thinking = call_api(client, fill_prompt, is_revision=True)
                                    
                                    if not fill_text.strip():
                                        st.warning("Fill pass returned empty.")
                                    else:
                                        if fill_thinking:
                                            with st.expander("Fill pass thinking"):
                                                st.text(fill_thinking[:3000])
                                        
                                        fill_word_count = len(fill_text.split())
                                        st.info(f"Fill pass: {cut_word_count:,} → {fill_word_count:,} words ({fill_word_count - cut_word_count:,} added)")
                                        
                                        fill_score = score_chapter(fill_text)
                                        
                                        # Report
                                        fill_report = generate_report(fill_score, fill_text, "After Fill", 2, st.session_state.source_profile)
                                        st.session_state.reports.append({"label": "After Fill", "buffer": fill_report})
                                        st.session_state.revision_history.append({"pass": 2, "text": fill_text, "label": "After Fill"})
                                        
                                        # Comparison
                                        rcol1, rcol2, rcol3 = st.columns(3)
                                        rcol1.metric("Red metrics", fill_score["red_count"], 
                                                   delta=f"{fill_score['red_count'] - cut_score['red_count']}", delta_color="inverse")
                                        rcol2.metric("Flagged sentences", len(fill_score["flagged"]), 
                                                   delta=f"{len(fill_score['flagged']) - len(cut_score['flagged'])}", delta_color="inverse")
                                        rcol3.metric("Final words", f"{fill_word_count:,}")
                                        
                                        display_scorecard(fill_score, st.session_state.source_profile)
                                        
                                        with st.expander("Filled Chapter Text", expanded=False):
                                            st.text(fill_text)
                                        
                                        fcol1, fcol2 = st.columns(2)
                                        with fcol1:
                                            fill_buf = make_docx(fill_text)
                                            st.download_button(
                                                label="Download Final (.docx)",
                                                data=fill_buf,
                                                file_name="chapter_final.docx",
                                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                key="ch_dl_fill"
                                            )
                                        with fcol2:
                                            st.download_button(
                                                label="Download Final Report (.docx)",
                                                data=st.session_state.reports[-1]["buffer"],
                                                file_name="report_final.docx",
                                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                key="report_dl_fill"
                                            )
                                        
                                        current_text = fill_text
                                        current_score = fill_score
                                
                                elif run_fill_pass and deficit <= 100:
                                    st.info(f"Only {deficit} words removed — too few to justify a fill pass.")
                    
                    # Update final state
                    st.session_state.chapter_text = current_text
                    st.session_state.score_result = current_score
                    
                    # Final summary
                    if len(st.session_state.revision_history) > 1:
                        st.markdown("---")
                        st.markdown("### Pipeline Summary")
                        orig_wc = original_word_count
                        final_wc = len(current_text.split())
                        orig_red = score_result["red_count"]
                        final_red = current_score["red_count"]
                        orig_flagged = len(score_result["flagged"])
                        final_flagged = len(current_score["flagged"])
                        
                        scol1, scol2, scol3, scol4 = st.columns(4)
                        scol1.metric("Words", f"{final_wc:,}", delta=f"{final_wc - orig_wc:,}")
                        scol2.metric("Red metrics", final_red, delta=f"{final_red - orig_red}", delta_color="inverse")
                        scol3.metric("Flagged sentences", final_flagged, delta=f"{final_flagged - orig_flagged}", delta_color="inverse")
                        scol4.metric("Overall", current_score["overall"])
            
            except anthropic.APIError as e:
                st.error(f"API Error: {e}")
            except Exception as e:
                st.error(f"Error: {e}")
                import traceback
                st.text(traceback.format_exc())


# ──────────────────────────────────────────────
# MANUAL SCORE & REVISE (for pasted text)
# ──────────────────────────────────────────────

st.markdown("---")
st.subheader("Score Existing Text")
st.caption("Paste any chapter text below to score it without writing a new one.")

pasted_text = st.text_area("Paste chapter text", height=200, key="paste_score")

if st.button("Score This Text"):
    if pasted_text.strip():
        score_result = score_chapter(pasted_text)
        display_scorecard(score_result, st.session_state.source_profile)
        st.session_state.chapter_text = pasted_text
        st.session_state.score_result = score_result
        
        report_buf = generate_report(score_result, pasted_text, "Pasted Text", 0, st.session_state.source_profile)
        st.download_button(
            label="Download Score Report (.docx)",
            data=report_buf,
            file_name="report_pasted.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="report_pasted"
        )
    else:
        st.warning("Paste some text first.")

if st.button("Cut This Text"):
    if not api_key:
        st.error("Enter your API key in the sidebar.")
    elif not pasted_text.strip():
        st.warning("Paste some text first.")
    else:
        score_result = score_chapter(pasted_text)
        if not score_result["flagged"]:
            st.info("No flagged passages to cut.")
        else:
            display_scorecard(score_result, st.session_state.source_profile)
            cutting_prompt = build_cutting_prompt(pasted_text, score_result)
            
            with st.spinner("Cutting..."):
                try:
                    client = anthropic.Anthropic(api_key=api_key)
                    cut_text, _ = call_api(client, cutting_prompt, is_revision=True)
                    
                    if cut_text.strip():
                        new_score = score_chapter(cut_text)
                        
                        orig_wc = len(pasted_text.split())
                        cut_wc = len(cut_text.split())
                        
                        st.markdown("### Cut Version")
                        st.info(f"{orig_wc:,} → {cut_wc:,} words ({orig_wc - cut_wc:,} removed)")
                        display_scorecard(new_score, st.session_state.source_profile)
                        
                        with st.expander("Cut text", expanded=False):
                            st.text(cut_text)
                        
                        rcol1, rcol2 = st.columns(2)
                        with rcol1:
                            buffer = make_docx(cut_text)
                            st.download_button(
                                label="Download Cut Version (.docx)",
                                data=buffer,
                                file_name="chapter_cut.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        with rcol2:
                            report_buf = generate_report(new_score, cut_text, "Cut (Pasted)", 1, st.session_state.source_profile)
                            st.download_button(
                                label="Download Cut Report (.docx)",
                                data=report_buf,
                                file_name="report_cut.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="report_cut_pasted"
                            )
                except Exception as e:
                    st.error(f"Error: {e}")


# ──────────────────────────────────────────────
# VERSION HISTORY
# ──────────────────────────────────────────────

if st.session_state.revision_history and len(st.session_state.revision_history) > 1:
    st.markdown("---")
    st.subheader("Version History")
    
    for entry in st.session_state.revision_history:
        with st.expander(f"{entry['label']} ({len(entry['text'].split()):,} words)"):
            st.text(entry["text"][:2000] + ("..." if len(entry["text"]) > 2000 else ""))
            dcol1, dcol2 = st.columns(2)
            with dcol1:
                buffer = make_docx(entry["text"])
                st.download_button(
                    label=f"Download {entry['label']}",
                    data=buffer,
                    file_name=f"chapter_{entry['label'].lower().replace(' ', '_').replace('(', '').replace(')', '')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_{entry['pass']}"
                )
            with dcol2:
                matching_reports = [r for r in st.session_state.reports if r["label"] == entry["label"]]
                if matching_reports:
                    st.download_button(
                        label=f"Download {entry['label']} Report",
                        data=matching_reports[0]["buffer"],
                        file_name=f"report_{entry['label'].lower().replace(' ', '_').replace('(', '').replace(')', '')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"rdl_{entry['pass']}"
                    )
